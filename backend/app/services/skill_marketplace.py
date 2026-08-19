"""Catálogo local de skills y tool de búsqueda para el agente."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from langchain_core.tools import tool

_SKILLS_DIR = Path(__file__).resolve().parents[1] / "skills"


def _load_skill_code(filename: str) -> str:
    return (_SKILLS_DIR / filename).read_text(encoding="utf-8")

SkillRecord = dict[str, Any]

_SKILL_CAUDAL = '''
def run(input_data):
    """Q = A * v. Área en m², velocidad en m/s → caudal en m³/s y L/s."""
    area_m2 = float(input_data.get("area_m2") or input_data.get("area") or 0)
    velocidad_ms = float(
        input_data.get("velocidad_ms")
        or input_data.get("velocidad")
        or input_data.get("v")
        or 0
    )
    caudal_m3s = area_m2 * velocidad_ms
    return {
        "area_m2": area_m2,
        "velocidad_ms": velocidad_ms,
        "caudal_m3s": round(caudal_m3s, 6),
        "caudal_ls": round(caudal_m3s * 1000, 3),
        "formula": "Q = A * v",
    }
'''

_SKILL_CONVERSION = '''
def run(input_data):
    """Convierte caudales entre L/s, m³/s, m³/h y m³/día."""
    value = float(input_data.get("valor") or input_data.get("value") or 0)
    unidad = str(
        input_data.get("unidad")
        or input_data.get("from")
        or input_data.get("de")
        or "l/s"
    ).lower().replace(" ", "")
    to_m3s = {
        "l/s": 0.001,
        "ls": 0.001,
        "lps": 0.001,
        "m3/s": 1.0,
        "m3s": 1.0,
        "m3/h": 1.0 / 3600.0,
        "m3h": 1.0 / 3600.0,
        "m3/d": 1.0 / 86400.0,
        "m3/dia": 1.0 / 86400.0,
        "m3d": 1.0 / 86400.0,
    }
    factor = to_m3s.get(unidad)
    if factor is None:
        return {"error": f"Unidad no soportada: {unidad}", "soportadas": list(to_m3s)}
    m3s = value * factor
    return {
        "entrada": {"valor": value, "unidad": unidad},
        "m3_s": round(m3s, 8),
        "l_s": round(m3s * 1000, 6),
        "m3_h": round(m3s * 3600, 6),
        "m3_dia": round(m3s * 86400, 4),
    }
'''

_SKILL_PRORRATEO = '''
def run(input_data):
    """Prorratea un caudal o volumen entre derechos (acciones / hectáreas)."""
    total = float(input_data.get("total") or input_data.get("caudal_total") or 0)
    partes = input_data.get("partes") or input_data.get("derechos") or []
    if isinstance(partes, dict):
        partes = [{"nombre": k, "peso": v} for k, v in partes.items()]
    pesos = []
    for item in partes:
        if isinstance(item, dict):
            pesos.append(
                {
                    "nombre": str(item.get("nombre") or item.get("id") or "parte"),
                    "peso": float(item.get("peso") or item.get("acciones") or item.get("ha") or 0),
                }
            )
        else:
            pesos.append({"nombre": str(item), "peso": 1.0})
    suma = sum(p["peso"] for p in pesos) or 1.0
    asignaciones = [
        {
            "nombre": p["nombre"],
            "peso": p["peso"],
            "proporcion": round(p["peso"] / suma, 6),
            "asignado": round(total * p["peso"] / suma, 6),
        }
        for p in pesos
    ]
    return {"total": total, "suma_pesos": suma, "asignaciones": asignaciones}
'''

_SKILL_LAMINA = '''
def run(input_data):
    """Lámina de riego (mm) a partir de volumen (m³) y superficie (ha o m²)."""
    volumen_m3 = float(input_data.get("volumen_m3") or input_data.get("volumen") or 0)
    superficie_ha = input_data.get("superficie_ha") or input_data.get("ha")
    superficie_m2 = input_data.get("superficie_m2") or input_data.get("m2")
    if superficie_ha is not None:
        area_m2 = float(superficie_ha) * 10000.0
        ha = float(superficie_ha)
    else:
        area_m2 = float(superficie_m2 or 0)
        ha = area_m2 / 10000.0 if area_m2 else 0.0
    lamina_mm = (volumen_m3 / area_m2) * 1000.0 if area_m2 else 0.0
    return {
        "volumen_m3": volumen_m3,
        "superficie_m2": area_m2,
        "superficie_ha": round(ha, 4),
        "lamina_mm": round(lamina_mm, 3),
        "formula": "lámina_mm = (volumen_m3 / área_m2) * 1000",
    }
'''

_SKILL_TIEMPO = '''
def run(input_data):
    """Tiempo de riego a partir de volumen (m³) y caudal (L/s o m³/s)."""
    volumen_m3 = float(input_data.get("volumen_m3") or input_data.get("volumen") or 0)
    caudal_ls = input_data.get("caudal_ls") or input_data.get("l_s")
    caudal_m3s = input_data.get("caudal_m3s") or input_data.get("m3_s")
    if caudal_ls is not None:
        q_m3s = float(caudal_ls) / 1000.0
    else:
        q_m3s = float(caudal_m3s or 0)
    segundos = volumen_m3 / q_m3s if q_m3s else 0.0
    horas = segundos / 3600.0
    return {
        "volumen_m3": volumen_m3,
        "caudal_m3s": round(q_m3s, 8),
        "caudal_ls": round(q_m3s * 1000, 4),
        "tiempo_s": round(segundos, 1),
        "tiempo_h": round(horas, 4),
        "tiempo_hm": {
            "horas": int(horas),
            "minutos": int(round((horas - int(horas)) * 60)),
        },
    }
'''

_SKILL_DOCX = '''
def run(input_data):
    """Genera un .docx con formato: headings, párrafos, viñetas, numeración y tablas."""
    import base64
    import re
    from io import BytesIO
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title = str(input_data.get("titulo") or input_data.get("title") or "Documento Irrigación")
    filename = str(input_data.get("filename") or input_data.get("nombre") or f"{title}.docx")
    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"

    def _add_formatted_runs(paragraph, text):
        """Soporta **negrita** y *cursiva* (no anidados)."""
        raw = str(text or "")
        if not raw:
            return
        pattern = re.compile(r"(\\*\\*[^*]+\\*\\*|\\*[^*]+\\*|[^*]+)")
        for token in pattern.findall(raw):
            if token.startswith("**") and token.endswith("**") and len(token) > 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*") and len(token) > 2:
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                paragraph.add_run(token)

    def _align(paragraph, value):
        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "izquierda": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "centro": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "derecha": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "justificado": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        key = str(value or "").strip().lower()
        if key in mapping:
            paragraph.alignment = mapping[key]

    def _list_items(block, text):
        items = block.get("items") or block.get("elementos")
        if isinstance(items, list):
            return items
        return [text] if text else []

    blocks = input_data.get("blocks") or input_data.get("bloques")
    if not isinstance(blocks, list):
        blocks = []

    if not blocks:
        content = str(
            input_data.get("contenido")
            or input_data.get("content")
            or input_data.get("texto")
            or ""
        ).strip()
        paragraphs = input_data.get("paragraphs") or input_data.get("parrafos")
        if isinstance(paragraphs, list) and paragraphs:
            blocks = [{"type": "paragraph", "text": str(p)} for p in paragraphs if str(p).strip()]
        elif content:
            parts = re.split(r"\\n\\s*\\n", content.replace("\\r\\n", "\\n"))
            blocks = [{"type": "paragraph", "text": p.strip()} for p in parts if p.strip()]

    if not blocks:
        blocks = [{"type": "paragraph", "text": "(Sin contenido)"}]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    include_title = bool(input_data.get("include_title", True))
    if include_title and title:
        h = doc.add_heading(title, level=0)
        _align(h, input_data.get("title_align") or "center")

    rendered = 0
    for block in blocks:
        if not isinstance(block, dict):
            text = str(block).strip()
            if text:
                p = doc.add_paragraph()
                _add_formatted_runs(p, text)
                rendered += 1
            continue

        tipo = str(block.get("type") or block.get("tipo") or "paragraph").lower().strip()
        text = block.get("text") or block.get("texto") or ""
        level = int(block.get("level") or block.get("nivel") or 1)
        level = max(1, min(level, 3))

        if tipo in {"heading", "titulo", "heading1", "heading2", "heading3", "h1", "h2", "h3"}:
            if tipo in {"heading1", "h1"}:
                level = 1
            elif tipo in {"heading2", "h2"}:
                level = 2
            elif tipo in {"heading3", "h3"}:
                level = 3
            h = doc.add_heading(str(text), level=level)
            _align(h, block.get("align") or block.get("alineacion"))
            rendered += 1
        elif tipo in {"bullet", "viñeta", "vineta", "ul", "lista"}:
            for item in _list_items(block, text):
                item_s = str(item).strip()
                if not item_s:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                _add_formatted_runs(p, item_s)
                rendered += 1
        elif tipo in {"number", "numbered", "ol", "numerada", "numerado"}:
            for item in _list_items(block, text):
                item_s = str(item).strip()
                if not item_s:
                    continue
                p = doc.add_paragraph(style="List Number")
                _add_formatted_runs(p, item_s)
                rendered += 1
        elif tipo in {"table", "tabla"}:
            rows = block.get("rows") or block.get("filas") or []
            if isinstance(rows, list) and rows:
                cols = max(len(r) if isinstance(r, (list, tuple)) else 1 for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    cells = list(row) if isinstance(row, (list, tuple)) else [row]
                    for j in range(cols):
                        cell_text = str(cells[j]) if j < len(cells) else ""
                        table.rows[i].cells[j].text = cell_text
                rendered += 1
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, str(text))
            _align(p, block.get("align") or block.get("alineacion"))
            rendered += 1

    buf = BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    return {
        "filename": filename,
        "content_base64": base64.b64encode(raw).decode("ascii"),
        "size_bytes": len(raw),
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "titulo": title,
        "block_count": rendered,
    }
'''

CATALOG: list[SkillRecord] = [
    {
        "id": "caudal_canal",
        "name": "Cálculo de caudal (Q = A·v)",
        "description": (
            "Calcula el caudal de un canal o toma a partir del área de la sección "
            "(m²) y la velocidad (m/s). Devuelve m³/s y L/s."
        ),
        "tags": ["caudal", "canal", "toma", "seccion", "velocidad", "q", "a", "v"],
        "code": _SKILL_CAUDAL.strip(),
    },
    {
        "id": "conversion_unidades",
        "name": "Conversión de unidades de caudal",
        "description": (
            "Convierte un caudal entre L/s, m³/s, m³/h y m³/día."
        ),
        "tags": ["conversion", "unidades", "l/s", "m3", "caudal", "equivalencia"],
        "code": _SKILL_CONVERSION.strip(),
    },
    {
        "id": "prorrateo_turno",
        "name": "Prorrateo de turno de riego",
        "description": (
            "Reparte un caudal o volumen total entre derechos, acciones o hectáreas."
        ),
        "tags": ["prorrateo", "turno", "reparto", "acciones", "derechos", "hectareas"],
        "code": _SKILL_PRORRATEO.strip(),
    },
    {
        "id": "lamina_riego",
        "name": "Lámina de riego",
        "description": (
            "Calcula la lámina aplicada en mm a partir del volumen (m³) y la superficie "
            "(ha o m²)."
        ),
        "tags": ["lamina", "riego", "milimetros", "volumen", "superficie", "hectareas"],
        "code": _SKILL_LAMINA.strip(),
    },
    {
        "id": "tiempo_riego",
        "name": "Tiempo de riego",
        "description": (
            "Calcula cuánto tiempo hay que regar para aplicar un volumen dado un caudal."
        ),
        "tags": ["tiempo", "duracion", "riego", "horas", "caudal", "volumen"],
        "code": _SKILL_TIEMPO.strip(),
    },
    {
        "id": "generar_documento_word",
        "name": "Generación de documento Word (.docx)",
        "description": (
            "Redacta y exporta un documento Word (.docx) respetando el formato pedido: "
            "títulos/subtítulos, párrafos, viñetas, listas numeradas, tablas y negritas/cursivas."
        ),
        "tags": [
            "word",
            "docx",
            "documento",
            "informe",
            "redactar",
            "exportar",
            "escribir",
            "archivo",
            "formato",
            "tabla",
            "viñetas",
            "lista",
            "titulo",
            "negrita",
        ],
        "code": _SKILL_DOCX.strip(),
    },
    {
        "id": "crawl_domain_links",
        "name": "Crawler de subenlaces del sitio",
        "description": (
            "Escanea un sitio web y extrae todos los subenlaces válidos y páginas "
            "internas (HTML y PDFs) descartando enlaces rotos 404."
        ),
        "tags": [
            "crawler",
            "crawl",
            "escanear",
            "escanea",
            "subenlaces",
            "enlaces",
            "links",
            "href",
            "sitio",
            "mapear",
            "scraping",
            "pdf",
            "normativa",
        ],
        "code": _load_skill_code("crawl_domain_links.py"),
    },
]


APPROVAL_KIND_EXECUTE = "execute_local"
APPROVAL_KIND_DOWNLOAD = "download_remote"

_SKILL_KEYWORDS: dict[str, tuple[str, ...]] = {
    "caudal_canal": (
        "caudal",
        "q=",
        "q =",
        "area",
        "área",
        "velocidad",
        "seccion",
        "sección",
        "canal",
        "toma",
    ),
    "conversion_unidades": (
        "convert",
        "conversion",
        "conversión",
        "equival",
        "l/s",
        "m3/s",
        "m³/s",
        "m3/h",
        "m3/d",
        "pasame a",
        "pasá a",
        "pasa a",
        "unidades",
    ),
    "prorrateo_turno": (
        "prorrate",
        "repart",
        "reparto",
        "turno",
        "derechos",
        "acciones",
        "hectarea",
        "hectárea",
    ),
    "lamina_riego": ("lamina", "lámina", "milimetros", "milímetros", "mm de riego"),
    "tiempo_riego": (
        "tiempo de riego",
        "cuanto regar",
        "cuánto regar",
        "duracion",
        "duración",
        "horas de riego",
    ),
    "generar_documento_word": (
        "word",
        "docx",
        "documento",
        "informe",
        "redact",
        "exportar",
        "generar archivo",
        "armame",
        "armá",
        "escribime",
        "escribí",
    ),
    "crawl_domain_links": (
        "escanear",
        "escanea",
        "crawler",
        "crawl",
        "subenlaces",
        "subenlace",
        "mapear sitio",
        "mapear el sitio",
        "buscar subenlaces",
        "enlaces válidos",
        "enlaces validos",
        "href",
    ),
}

_CALC_VERBS = (
    "calcul",
    "convert",
    "prorrate",
    "equival",
    "repart",
    "cuánto da",
    "cuanto da",
    "pasame a",
    "pasá a",
    "pasa a",
    "cuanto es",
    "cuánto es",
    "dame el",
    "decime cu",
)

_UNCERTAINTY_MARKERS = (
    "no puedo calcular",
    "no tengo una herramienta",
    "no dispongo de",
    "no puedo hacer ese cálculo",
    "necesitaría datos numéricos",
    "faltan datos",
    "no encontré una skill",
    "reformulá con números",
    # Negativas de capacidad: deben disparar búsqueda/descarga de skill.
    "no puedo acceder",
    "no puedo entrar",
    "no puedo visitar",
    "páginas web externas",
    "paginas web externas",
    "en tiempo real",
    "mi capacidad para acceder",
    "no tengo acceso a internet",
    "no puedo navegar",
    "no puedo consultar sitios",
    "fuera de mi alcance",
    "no estoy habilitado",
    "no puedo realizar esa acción",
    "no puedo hacer eso",
)

_WEB_OR_EXTERNAL_HINTS = (
    "http://",
    "https://",
    "www.",
    "entrar a",
    "acceder a",
    "abrí la",
    "abri la",
    "abrir la página",
    "abrir la pagina",
    "consultar la web",
    "consultar la página",
    "consultar la pagina",
    "scrap",
    "telemetr",
    "serviciosweb",
    "desde internet",
    "en la web",
    # Ojo: NO usar "en la página" suelto — suele ser comentario del hilo
    # ("en la página dice 4cm"), no un pedido nuevo de scrapear.
    "altura en el punto",
    "altura del punto",
    "dato del punto",
)


def contains_url(text: str) -> bool:
    return bool(re.search(r"https?://", text or "", re.I))


def looks_like_web_or_external_request(text: str) -> bool:
    lowered = (text or "").lower()
    if not lowered.strip():
        return False
    # Réplicas al hilo ("en la página dice X") no son un pedido web nuevo.
    if is_result_challenge_or_correction(text):
        return False
    if is_site_crawl_request(text):
        return False
    if contains_url(lowered):
        return True
    return any(hint in lowered for hint in _WEB_OR_EXTERNAL_HINTS)


_CRAWL_INTENT_RE = re.compile(
    r"(?:"
    r"escan(?:ear|[eé]a)\s+(?:todos\s+los\s+)?(?:sub)?enlaces"
    r"|(?:sub)?enlaces\s+v[aá]lidos"
    r"|crawler?"
    r"|crawl(?:ear)?"
    r"|mapear\s+(?:el\s+)?sitio"
    r"|buscar\s+subenlaces"
    r"|listar\s+(?:los\s+)?(?:links|enlaces|href)"
    r")",
    re.I,
)

_SKILL_SEARCH_INTENT_RE = re.compile(
    r"(?:buscar|crear|generar|descargar|buscar/crear)\s+(?:una\s+)?skill"
    r"|search_skill_marketplace",
    re.I,
)

_TELEMETRY_INTENT_RE = re.compile(
    r"\b(?:"
    r"telemetr[ií]a"
    r"|sensor(?:es)?"
    r"|estaci[oó]n(?:es)?"
    r"|nivel\s+de\s+r[ií]o"
    r"|dato-medicions"
    r"|fulldto"
    r"|altura\s+(?:del|en\s+el)\s+punto"
    r"|caudal\s+(?:del|en\s+el)\s+punto"
    r")\b",
    re.I,
)


def is_site_crawl_request(text: str) -> bool:
    """Pedido de mapear/escanear un sitio; no es consulta de telemetría."""
    blob = (text or "").strip()
    if not blob:
        return False
    return bool(_CRAWL_INTENT_RE.search(blob))


def wants_skill_marketplace(text: str) -> bool:
    return bool(_SKILL_SEARCH_INTENT_RE.search(text or ""))


def is_result_challenge_or_correction(text: str) -> bool:
    """El usuario cuestiona/corrige un resultado previo del mismo hilo."""
    lowered = (text or "").lower().strip()
    if not lowered:
        return False
    patterns = (
        r"de\s+d[oó]nde\s+sacaste",
        r"de\s+donde\s+sali[oó]",
        r"c[oó]mo\s+calculaste",
        r"est[aá]\s+mal",
        r"no\s+es\s+(?:eso|correcto|as[ií]|la\s+altura)",
        r"(?:en\s+la\s+p[aá]gina|ah[ií]|ac[aá])\s+dice",
        r"dice\s+\d",
        r"inventaste",
        r"alucin",
        r"no\s+coincide",
        r"fijate\s+que",
        r"mir[aá]\s+que\s+(?:en|la|el|dice)",
        r"pero\s+en\s+la\s+p[aá]gina",
        r"sacaste\s+mal",
        r"te\s+equivocaste",
        r"te\s+dije\s+que",
        r"te\s+ped[ií]\s+que",
        r"no\s+era\s+ahora",
        r"no\s+lo\s+mandes\s+ahora",
        r"ten[ií]a\s+que\s+ser\s+m[aá]s\s+tarde",
    )
    return any(re.search(p, lowered) for p in patterns)


_DOMAIN_CHAT_RE = re.compile(
    r"\b(?:altura|caudal|padr[oó]n|expediente|norma|l[aá]mina|"
    r"prorrateo|telemetr|punto\s+\d|m3|l/s|hect[aá]rea|turno\s+de\s+riego)\b",
    re.I,
)


def is_casual_chat(
    text: str,
    *,
    history: list[dict[str, Any]] | None = None,
    context_text: str | None = None,
    thread_state: dict[str, Any] | None = None,
) -> bool:
    """Charla (saludo, cómo andás). No es orden, ni dato de riego, ni follow-up."""
    from app.services.context_memory import looks_like_save_context_intent
    from app.services.google_assistant import detect_google_intent
    from app.services.order_parse import looks_like_do_task

    blob = (text or "").strip()
    if not blob:
        return False
    if looks_like_do_task(blob) or looks_like_save_context_intent(blob):
        return False
    if detect_google_intent(blob) is not None:
        return False
    if should_try_skill_marketplace(blob):
        return False
    if is_site_crawl_request(blob) or wants_skill_marketplace(blob):
        return False
    if looks_like_web_or_external_request(blob) or contains_url(blob):
        return False
    if is_asking_for_needed_data(blob) or is_result_challenge_or_correction(blob):
        return False
    if is_download_confirmation_only(blob):
        return False
    ctx = context_text or conversation_context_text(blob, history)
    if is_thread_followup(
        blob, ctx, thread_state, history=history
    ) and extract_open_task(
        blob, history, ctx, thread_state=thread_state
    ):
        return False
    if _DOMAIN_CHAT_RE.search(blob):
        return False
    if len(blob) > 160 and "?" in blob:
        return False
    return True


_SWITCH_STOP = {
    "para",
    "como",
    "qué",
    "que",
    "una",
    "unos",
    "unas",
    "los",
    "las",
    "del",
    "con",
    "por",
    "me",
    "te",
    "se",
    "le",
    "de",
    "la",
    "el",
    "en",
    "un",
    "al",
    "lo",
    "y",
    "o",
    "a",
    "hace",
    "hacé",
    "hacer",
    "podes",
    "podés",
    "podrias",
    "podrías",
    "necesito",
    "quiero",
    "escuchame",
    "mira",
    "che",
    "bld",
    "crack",
    "todo",
    "bien",
    "orden",
}


def _content_tokens(text: str) -> set[str]:
    return {t for t in _tokenize(text) if t not in _SWITCH_STOP and len(t) > 2}


def previous_open_task(
    *,
    context_text: str | None = None,
    history: list[dict[str, Any]] | None = None,
    thread_state: dict[str, Any] | None = None,
    current_text: str | None = None,
) -> str:
    """Tarea YA abierta (sin pisar con el mensaje actual)."""
    from app.services.thread_memory import open_task_from_state

    persisted = open_task_from_state(thread_state)
    if persisted:
        return persisted[:300]
    ctx = (context_text or "").strip()
    current = (current_text or "").strip()
    if current and ctx.endswith(current):
        ctx = ctx[: -len(current)].strip()
    return extract_open_task("", history, ctx or None, thread_state=None)[:300]


def is_context_switch(
    text: str,
    *,
    context_text: str | None = None,
    history: list[dict[str, Any]] | None = None,
    thread_state: dict[str, Any] | None = None,
) -> bool:
    """Pedido nuevo autónomo: no es dato ni respuesta de la tarea abierta."""
    from app.services.context_memory import looks_like_save_context_intent, parse_context_scope
    from app.services.google_assistant import detect_google_intent
    from app.services.order_parse import looks_like_do_task

    blob = (text or "").strip()
    if not blob:
        return False
    if is_download_confirmation_only(blob) or is_asking_for_needed_data(blob):
        return False
    if is_result_challenge_or_correction(blob) or parse_context_scope(blob):
        return False
    previous = previous_open_task(
        context_text=context_text,
        history=history,
        thread_state=thread_state,
        current_text=blob,
    )
    if not previous:
        return False
    # Dato suelto (URL, IP, mail) sin verbo de pedido nuevo → sigue el hilo.
    if (
        (has_concrete_task_inputs(blob) or contains_url(blob))
        and not looks_like_do_task(blob)
        and not is_action_request(blob)
    ):
        return False
    autonomous = (
        looks_like_do_task(blob)
        or looks_like_save_context_intent(blob)
        or detect_google_intent(blob) is not None
        or should_try_skill_marketplace(blob)
        or is_action_request(blob)
    )
    if not autonomous:
        return False
    prev_l = previous.lower().strip()
    new_l = blob.lower()
    if prev_l and (prev_l in new_l or new_l in prev_l):
        return False
    prev_toks = _content_tokens(previous)
    new_toks = _content_tokens(blob)
    if prev_toks and new_toks:
        overlap = len(prev_toks & new_toks) / len(new_toks)
        if overlap >= 0.5:
            return False
    return True


def is_thread_followup(
    text: str,
    context_text: str | None = None,
    thread_state: dict[str, Any] | None = None,
    history: list[dict[str, Any]] | None = None,
) -> bool:
    """Mensaje que continúa la tarea ya abierta (cualquier tema), no un pedido nuevo."""
    if is_context_switch(
        text,
        context_text=context_text,
        history=history,
        thread_state=thread_state,
    ):
        return False
    if is_download_confirmation_only(text):
        return True
    if is_result_challenge_or_correction(text):
        return True
    if is_asking_for_needed_data(text):
        ctx = (context_text or "").strip()
        from app.services.thread_memory import open_task_from_state

        return bool(
            (ctx and ctx != (text or "").strip()) or open_task_from_state(thread_state)
        )
    ctx = (context_text or "").strip()
    if not ctx or ctx == (text or "").strip():
        from app.services.thread_memory import open_task_from_state

        if not open_task_from_state(thread_state):
            return False
    cleaned = (text or "").strip()
    if not cleaned:
        return False
    open_task = previous_open_task(
        context_text=ctx,
        history=history,
        thread_state=thread_state,
        current_text=cleaned,
    )
    if not open_task:
        return False
    from app.services.order_parse import looks_like_do_task

    if looks_like_do_task(cleaned) and not is_asking_for_needed_data(cleaned):
        same = open_task.lower() in cleaned.lower() or cleaned.lower() in open_task.lower()
        if not same:
            return False
    if contains_url(cleaned) and len(cleaned) < 220:
        return True
    args = extract_web_skill_args(cleaned)
    if args.get("punto") and len(_tokenize(cleaned)) <= 10:
        return True
    if (
        not looks_like_do_task(cleaned)
        and not is_action_request(cleaned)
        and len(cleaned) <= 200
        and len(_tokenize(cleaned)) <= 18
    ):
        return True
    return False


_DATA_ASK_RE = re.compile(
    r"(?:qu[eé]\s+datos|qu[eé]\s+te\s+falta|qu[eé]\s+necesit(?:[aá]|as|o)|"
    r"faltan\s+datos|con\s+qu[eé]\s+datos|qu[eé]\s+me\s+ten[eé]s\s+que\s+pasar|"
    r"qu[eé]\s+info(?:rmaci[oó]n)?)",
    re.I,
)


def is_asking_for_needed_data(text: str) -> bool:
    """El usuario pregunta qué datos hacen falta para la tarea ya abierta."""
    return bool(_DATA_ASK_RE.search(text or ""))


def has_concrete_task_inputs(text: str) -> bool:
    """Hay un dato usable (IP, rango, URL, MAC, email, magnitudes)."""
    blob = text or ""
    return bool(
        re.search(r"\b\d{1,3}(?:\.\d{1,3}){3}(?:/\d{1,2})?\b", blob)
        or re.search(r"(?:[0-9a-f]{2}:){5}[0-9a-f]{2}", blob, re.I)
        or contains_url(blob)
        or re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", blob)
        or re.search(
            r"\b\d+(?:[.,]\d+)?\s*(?:ha|m2|m²|l/s|m3|m³)",
            blob,
            re.I,
        )
    )


def extract_open_task(
    user_message: str = "",
    history: list[dict[str, Any]] | None = None,
    context_text: str | None = None,
    thread_state: dict[str, Any] | None = None,
) -> str:
    """Tarea activa del hilo (cualquier tema), no el último chiste ni un 'qué datos'."""
    from app.services.order_parse import looks_like_do_task
    from app.services.thread_memory import open_task_from_state

    history = history or []
    current = (user_message or "").strip()
    if (
        current
        and looks_like_do_task(current)
        and not is_asking_for_needed_data(current)
        and not is_download_confirmation_only(current)
    ):
        return re.sub(r"\s+", " ", current).strip()[:300]
    persisted = open_task_from_state(thread_state)
    if persisted:
        return persisted[:300]
    for item in reversed(history):
        role = (item.get("role") or "").lower()
        msg = (item.get("message") or "").strip()
        if role not in {"assistant", "ai"} or not msg:
            continue
        named = re.search(
            r"skill\s+['«\"]([^'»\"]+)['»\"]|"
            r"seguimos con\s+\*\*([^*]+)\*\*|"
            r"para (?:usar\s+)?\*\*([^*]+)\*\*",
            msg,
            re.I,
        )
        if named:
            captured = next((g.strip() for g in named.groups() if g), "")
            from app.services.llm_roles import sanitize_open_task

            captured = sanitize_open_task(captured)
            if captured:
                return captured[:300]
    user_msgs: list[str] = []
    for item in history:
        if (item.get("role") or "").lower() == "user":
            msg = (item.get("message") or "").strip()
            if msg:
                user_msgs.append(msg)
    if not user_msgs and context_text:
        user_msgs = [
            block.strip()
            for block in re.split(r"\n\s*\n", context_text)
            if block.strip()
        ]
    if current and (not user_msgs or user_msgs[-1] != current):
        user_msgs.append(current)
    for msg in reversed(user_msgs):
        if is_asking_for_needed_data(msg) or is_download_confirmation_only(msg):
            continue
        if looks_like_do_task(msg) or has_actionable_remote_task(msg):
            return re.sub(r"\s+", " ", msg).strip()[:300]
    return ""


def ask_inputs_for_open_task(
    context: str,
    skill_name: str | None = None,
    *,
    history: list[dict[str, Any]] | None = None,
    thread_state: dict[str, Any] | None = None,
) -> str:
    """Seguí la tarea ABIERTA del hilo, sea cual sea. No cambies de tema."""
    from app.services.llm_roles import sanitize_open_task

    task = sanitize_open_task(skill_name) or extract_open_task(
        "", history, context, thread_state=thread_state
    )
    task = sanitize_open_task(re.sub(r"\s+", " ", task).strip()[:220])
    if not task:
        return (
            "Decime qué necesitás hacer y te digo qué datos hacen falta. "
            "No arranco un trámite viejo si no hay uno abierto."
        )
    return (
        f"Seguimos con lo de **{task}**, no arranco otra cosa.\n\n"
        "Tirame los datos concretos que apliquen (destino, archivo, rango, "
        "URL, lo que sea) y lo hago. Si no sabés cuáles, mandá lo que tengas "
        "y te digo qué falta."
    )


def missing_web_clarify_fields(context_text: str) -> list[str]:
    """Qué falta preguntar para un pedido web, según el hilo completo."""
    probe = context_text or ""
    if is_site_crawl_request(probe):
        return [] if contains_url(probe) else ["url"]
    missing: list[str] = []
    args = extract_web_skill_args(probe)
    lowered = probe.lower()
    has_url = contains_url(probe) or is_telemetria_request(probe)
    if not is_telemetria_request(probe):
        if not has_url:
            missing.append("url")
        return missing
    has_target = bool(args.get("punto"))
    wants_metric = any(
        k in lowered
        for k in ("altura", "caudal", "nivel", "dato", "medición", "medicion", "valor", "telemetr")
    )
    if not has_url:
        missing.append("url")
    if not has_target:
        missing.append("punto")
    if not wants_metric:
        missing.append("metric")
    return missing


def thread_brief_for_prompt(
    user_message: str,
    history: list[dict[str, Any]] | None = None,
    thread_state: dict[str, Any] | None = None,
) -> str:
    """Resumen corto del hilo para que el LLM no trate cada mensaje como chat nuevo."""
    history = history or []
    effective = resolve_effective_remote_task(
        user_message,
        conversation_context_text(user_message, history),
    )
    last_assistant = ""
    for item in reversed(history):
        if (item.get("role") or "").lower() in {"assistant", "ai", "system"}:
            last_assistant = (item.get("message") or "").strip()
            if last_assistant:
                break
    args = extract_web_skill_args(effective)
    open_task = extract_open_task(
        user_message, history, thread_state=thread_state
    )
    persisted_text = str((thread_state or {}).get("summary_text") or "").strip()
    if is_context_switch(
        user_message, history=history, thread_state=thread_state
    ):
        return "\n".join(
            [
                "CAMBIO DE TAREA: el usuario empezó un pedido NUEVO.",
                f"- Tarea actual: {(open_task or user_message)[:500]}",
                "- No retomes la tarea anterior ni pidas sus datos "
                "(destino, archivo, URL) salvo que ESTE pedido los necesite.",
            ]
        )
    lines = [
        "CONTINUIDAD DEL HILO (obligatorio respetar):",
        "- Este mensaje es continuación del mismo chat, no un pedido aislado.",
        "- No vuelvas a pedir URL/punto/dato si ya están en el historial.",
        "- Si el usuario corrige un resultado, explicá la fuente o reconsultá; no reinicies el cuestionario.",
        "- Si pregunta qué datos faltan o aporta un dato, SEGUÍ la tarea abierta. "
        "Prohibido ofrecer el catálogo de riego u otra skill salvo que esa sea la tarea.",
    ]
    if open_task:
        lines.append(f"- TAREA ABIERTA (no la abandones): {open_task[:500]}")
    elif effective and effective.strip() != (user_message or "").strip():
        lines.append(f"- Tarea activa del hilo: {effective[:500]}")
    if args.get("url"):
        lines.append(f"- URL ya conocida: {args['url']}")
    if args.get("punto"):
        lines.append(f"- Punto ya conocido: {args['punto']}")
    if last_assistant:
        lines.append(f"- Última respuesta tuya (recorte): {last_assistant[:280]}")
    if is_result_challenge_or_correction(user_message):
        lines.append(
            "- El usuario cuestiona/corrige el resultado anterior: respondé sobre ese hilo."
        )
    if persisted_text:
        lines.append(persisted_text)
    return "\n".join(lines)


def reply_is_capability_refusal(reply: str | None) -> bool:
    if not reply:
        return False
    lowered = reply.lower()
    return any(marker in lowered for marker in _UNCERTAINTY_MARKERS)


def should_try_skill_marketplace(text: str, assistant_reply: str | None = None) -> bool:
    """Determina si conviene buscar en el marketplace antes de responder."""
    lowered = (text or "").lower()
    if not lowered.strip():
        return False
    if is_site_crawl_request(text) or wants_skill_marketplace(text):
        return True

    if re.search(
        r"\b(?:skills?|habilidades?|herramientas?|scripts?)\b.{0,40}"
        r"(?:para|de|que|buscar|busc|encontr|instal|descarg|gener)"
        r"|(?:buscar|busc(?:á|a|ar)|encontr(?:á|a|ar)|instal(?:á|a|ar)|"
        r"descarg(?:á|a|ar)|gener(?:á|a|ar)).{0,40}"
        r"\b(?:skills?|habilidades?|herramientas?|scripts?)\b",
        lowered,
    ):
        return True
    if looks_like_web_or_external_request(text):
        return True
    if any(kw in lowered for keywords in _SKILL_KEYWORDS.values() for kw in keywords):
        return True
    if any(verb in lowered for verb in _CALC_VERBS):
        return True
    if re.search(r"\d", lowered) and any(
        term in lowered
        for term in ("caudal", "riego", "convert", "prorrate", "lamina", "lámina", "m3", "l/s")
    ):
        return True

    if assistant_reply and reply_is_capability_refusal(assistant_reply):
        return True
    from app.services.order_parse import looks_like_do_task

    if looks_like_do_task(text):
        return True
    return False


_ACTION_REQUEST_PATTERNS = tuple(
    re.compile(pattern)
    for pattern in (
        r"\bpodr[ií]as?\b",
        r"\bpod[eé]s\b",
        r"\bpodes\b",
        r"\bpuede[s]?\b",
        r"\bsabr[ií]as?\b",
        r"\bsabr[eé]s\b",
        r"\bsabes\b",
        r"\bhac[eé]lo\b",
        r"\bhac[eé]\s",
        r"\bhace\s",
        r"\bquiero que\b",
        r"\bnecesito que\b",
        r"\bme pod[eé]s\b",
        r"\bme podes\b",
        r"\bme podr[ií]as?\b",
        r"\btien[eé]s\s+(?:alguna|una)\s+(?:skill|habilidad|herramienta)",
        r"\bpod[eé]s\s+(?:calcular|convertir|generar|exportar|redactar|armar|entrar|acceder|consultar)",
        r"\bentr(ar|á|a)\s+a\b",
        r"\bacced(er|é|e)\s+a\b",
        r"\bconsult(á|a|ar)\b",
    )
)


def is_action_request(text: str) -> bool:
    """Detecta pedidos del tipo 'hacé esto', web externa o '¿podés hacer esto?'."""
    lowered = (text or "").lower().strip()
    if not lowered:
        return False
    if is_site_crawl_request(text) or wants_skill_marketplace(text):
        return True
    if looks_like_web_or_external_request(text):
        return True
    if should_try_skill_marketplace(text):
        return True
    return any(pattern.search(lowered) for pattern in _ACTION_REQUEST_PATTERNS)


def looks_like_skill_intent(text: str) -> bool:
    """Heurística amplia: cálculos, conversiones, automatizaciones y documentos."""
    return is_action_request(text)


def match_catalog_by_keywords(text: str) -> dict[str, Any] | None:
    """Coincidencia directa de palabras clave del usuario contra skills locales."""
    ranked = rank_catalog_skills(text)
    if not ranked:
        return None
    best, best_score = ranked[0]
    if best_score < 4:
        return None
    if len(ranked) > 1 and best_score - ranked[1][1] < 2:
        return None
    return search_catalog(best["name"], infer_arguments(best["id"], text))


# Palabras que aparecen en muchas tareas (no identifican una skill de riego).
_WEAK_CATALOG_TOKENS = {
    "velocidad",
    "tiempo",
    "test",
    "valor",
    "dato",
    "datos",
    "area",
    "área",
    "hacer",
    "calculo",
    "calcular",
    "calcula",
    "unidades",
    "conversion",
    "conversión",
    "internet",
}


def rank_catalog_skills(task: str) -> list[tuple[SkillRecord, int]]:
    """Puntúa todas las skills del catálogo (mayor = más pertinente)."""
    tokens = _tokenize(task)
    # Filtrar tokens demasiado genéricos que ensucian el score.
    stop = {
        "para",
        "como",
        "qué",
        "que",
        "una",
        "unos",
        "unas",
        "los",
        "las",
        "del",
        "con",
        "por",
        "hay",
        "esta",
        "este",
        "esto",
        "desde",
        "sobre",
        "me",
        "te",
        "se",
        "le",
        "de",
        "la",
        "el",
        "en",
        "un",
        "al",
        "lo",
        "y",
        "o",
        "a",
    }
    tokens = {t for t in tokens if t not in stop and len(t) > 2}
    lowered_task = (task or "").lower()
    ranked: list[tuple[SkillRecord, int]] = []
    for skill in CATALOG:
        haystack = " ".join(
            [skill["id"], skill["name"], skill["description"], " ".join(skill["tags"])]
        ).lower()
        explained = {t for t in tokens if t in haystack}
        distinctive = explained - _WEAK_CATALOG_TOKENS
        coverage = (len(explained) / len(tokens)) if tokens else 0.0
        # Una palabra genérica compartida no es esa skill (ej. "velocidad" ≠ caudal).
        if not distinctive and coverage < 0.45:
            continue
        if coverage < 0.22 and not distinctive:
            continue
        score = sum(1 for token in tokens if token in haystack)
        if skill["id"] in lowered_task or skill["name"].lower() in lowered_task:
            score += 5
        for keyword in _SKILL_KEYWORDS.get(skill["id"], ()):
            if keyword in lowered_task:
                score += 1 if keyword.strip().lower() in _WEAK_CATALOG_TOKENS else 3
        # Tags exactos (palabra completa) pesan más que substring accidental.
        for tag in skill["tags"]:
            if len(tag) >= 4 and re.search(rf"\b{re.escape(tag)}\b", lowered_task):
                score += 1 if tag.lower() in _WEAK_CATALOG_TOKENS else 2
        if score > 0 and (distinctive or coverage >= 0.45):
            ranked.append((skill, score))
    ranked.sort(key=lambda item: item[1], reverse=True)
    return ranked


def find_local_skill(task: str, arguments: dict[str, Any] | None = None) -> dict[str, Any]:
    """Busca una skill instalada con match de alta confianza."""
    available = [
        {"id": s["id"], "name": s["name"], "description": s["description"]}
        for s in CATALOG
    ]
    if looks_like_web_or_external_request(task) and not is_site_crawl_request(task):
        return {"found": False, "query": task, "available": available, "reason": "web_request"}

    ranked = rank_catalog_skills(task)
    if not ranked:
        return {"found": False, "query": task, "available": available, "reason": "no_signal"}

    best, best_score = ranked[0]
    second_score = ranked[1][1] if len(ranked) > 1 else 0

    # Ambigüedad: dos skills competidoras sin margen claro.
    if best_score >= 3 and second_score >= 3 and (best_score - second_score) < 2:
        return {
            "found": False,
            "query": task,
            "available": available,
            "ambiguous": True,
            "candidates": [
                {
                    "id": s["id"],
                    "name": s["name"],
                    "description": s["description"],
                    "score": sc,
                }
                for s, sc in ranked[:3]
            ],
            "reason": "ambiguous",
        }

    # Umbral alto: evita matches flojos (ej. un número suelto).
    if best_score < 4:
        return {
            "found": False,
            "query": task,
            "available": available,
            "top_score": best_score,
            "reason": "low_confidence",
        }

    resolved_args = arguments or {}
    inferred = infer_arguments(best["id"], task)
    merged_args = {**inferred, **{k: v for k, v in resolved_args.items() if v not in (None, "", {})}}
    return {
        "found": True,
        "id": best["id"],
        "name": best["name"],
        "description": best["description"],
        "code": best["code"],
        "arguments": merged_args,
        "score": best_score,
        "confidence": "high" if best_score - second_score >= 3 else "medium",
    }


def missing_fields_for_skill(skill_id: str, args: dict[str, Any]) -> list[str]:
    """Campos requeridos que aún faltan (ids de schema)."""
    schema = SKILL_ARG_SCHEMAS.get(skill_id)
    if not schema:
        return []
    missing: list[str] = []
    for key in schema.get("required") or []:
        if key == "partes":
            partes = args.get("partes") or args.get("derechos")
            if not isinstance(partes, list) or not partes:
                missing.append(key)
            continue
        if key in {"base_url", "url"}:
            if not str(args.get("base_url") or args.get("url") or "").strip():
                missing.append("base_url")
            continue
        if not _numeric_ready(args.get(key)):
            missing.append(key)
    if skill_id == "lamina_riego":
        if not (
            _numeric_ready(args.get("superficie_ha"))
            or _numeric_ready(args.get("superficie_m2"))
        ):
            if "superficie_ha" not in missing and "superficie_m2" not in missing:
                missing.append("superficie_ha_o_m2")
    if skill_id == "tiempo_riego":
        if not (
            _numeric_ready(args.get("caudal_ls")) or _numeric_ready(args.get("caudal_m3s"))
        ):
            missing.append("caudal_ls_o_m3s")
    return missing


def skill_missing_required_inputs(
    skill: dict[str, Any],
    args: dict[str, Any] | None = None,
    *,
    context_text: str = "",
    user_message: str = "",
) -> list[str]:
    """Qué falta de verdad. Vacío = se puede ejecutar (zero-arg OK, ej. speedtest)."""
    skill = skill or {}
    args = args or {}
    skill_id = str(skill.get("id") or "")
    if skill_id == "generar_documento_word":
        return []
    if skill_id == "crawl_domain_links":
        probe = f"{context_text or ''}\n{user_message or ''}".strip()
        if not str(args.get("base_url") or args.get("url") or "").strip() and contains_url(
            probe
        ):
            from app.services.skill_http import extract_urls

            urls = extract_urls(probe)
            if urls:
                args["base_url"] = urls[0]
                args["url"] = urls[0]
    missing = missing_fields_for_skill(skill_id, args)
    if missing:
        return missing
    if skill_id in SKILL_ARG_SCHEMAS:
        return []
    probe = f"{context_text or ''}\n{user_message or ''}".strip()
    desc = str(skill.get("description") or "")
    if is_telemetria_request(probe) or is_telemetria_request(desc):
        web_missing = missing_web_clarify_fields(probe)
        if args.get("punto"):
            web_missing = [m for m in web_missing if m != "punto"]
        if args.get("url") or args.get("api_url") or contains_url(probe):
            web_missing = [m for m in web_missing if m != "url"]
        return web_missing
    return []


def clarifying_question_for_skill(
    skill: dict[str, Any],
    missing: list[str],
    *,
    ambiguous_candidates: list[dict[str, Any]] | None = None,
) -> str:
    """Pregunta concreta al usuario cuando falta claridad o datos."""
    if ambiguous_candidates:
        options = "\n".join(
            f"- **{c.get('name')}**: {c.get('description')}" for c in ambiguous_candidates[:4]
        )
        return (
            "No estoy 100% seguro de qué herramienta usar. "
            "¿Cuál de estas necesitás?\n\n"
            f"{options}\n\n"
            "Respondeme con el nombre de la que corresponde (o reformulá el pedido)."
        )

    skill_name = skill.get("name") or skill.get("id") or "la skill"
    schema = SKILL_ARG_SCHEMAS.get(str(skill.get("id") or ""), {})
    labels = dict(schema.get("fields") or {})
    label_map = {
        **labels,
        "superficie_ha_o_m2": "superficie (ha o m²)",
        "caudal_ls_o_m3s": "caudal (L/s o m³/s)",
        "partes": "partes/derechos con sus pesos (acciones o ha)",
    }
    needed = [label_map.get(m, m) for m in missing]
    if not needed:
        return (
            f"Para usar **{skill_name}** me faltan datos. "
            "¿Me pasás los valores con unidades?"
        )
    bullets = "\n".join(f"- {item}" for item in needed)
    return (
        f"Entiendo que querés usar **{skill_name}**, pero me faltan datos concretos:\n\n"
        f"{bullets}\n\n"
        "Pasámelos en un mensaje (con unidades) y lo calculo."
    )


def clarifying_question_for_unknown(
    user_message: str,
    *,
    context_text: str | None = None,
    thread_state: dict[str, Any] | None = None,
) -> str:
    """Cuando el pedido es acción pero no hay match claro ni conviene adivinar."""
    probe = (context_text or user_message or "").strip() or (user_message or "")
    open_task = extract_open_task(
        user_message or "", context_text=probe, thread_state=thread_state
    )
    if open_task and not looks_like_web_or_external_request(open_task):
        return ask_inputs_for_open_task(probe, thread_state=thread_state)
    if looks_like_web_or_external_request(open_task or probe) or looks_like_web_or_external_request(
        user_message
    ) or is_telemetria_request(open_task or probe):
        missing = missing_web_clarify_fields(probe)
        label = {
            "url": "¿Qué URL/API exacta? (pegá el link)",
            "punto": "¿Código de punto/estación? (ej. 10009)",
            "metric": "¿Qué dato necesitás? (altura, caudal, última medición…)",
        }
        if not missing:
            return (
                "Con lo que ya me pasaste en el chat alcanza. "
                "Si querés que genere/ejecute la skill, respondé **sí, descargá la skill**."
            )
        bullets = "\n".join(f"- {label[m]}" for m in missing if m in label)
        known = extract_web_skill_args(probe)
        known_bits = []
        if known.get("url"):
            known_bits.append(f"URL: {known['url']}")
        if known.get("punto"):
            known_bits.append(f"punto: {known['punto']}")
        known_line = (
            ("Ya tengo: " + "; ".join(known_bits) + ".\n\n") if known_bits else ""
        )
        return (
            f"{known_line}"
            "Para seguir con esa consulta web/API me falta solo esto:\n"
            f"{bullets}\n\n"
            "Respondé con eso (o **sí, descargá la skill** si ya está completo en el hilo)."
        )
    examples = ", ".join(s["name"] for s in CATALOG[:4])
    return (
        "No terminé de entender qué automatización necesitás. "
        "¿Podés decirme en una frase qué resultado querés y con qué datos?\n\n"
        f"Ejemplos de lo que sí puedo hacer del catálogo: {examples}, "
        "o generar un Word / consultar una web institucional.\n\n"
        "Si querés que arme una skill nueva para algo fuera del catálogo, "
        "decime **descargá una skill** y describí la tarea."
    )


def resolve_skill_decision(
    user_message: str,
    arguments: dict[str, Any] | None = None,
    *,
    context_text: str | None = None,
    thread_state: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """
    Decisión endurecida de skills.

    action:
      - execute: skill local clara (puede faltar HITL afuera)
      - download: pedido web/externo o acción clara sin skill local
      - clarify: falta certeza o faltan argumentos
      - none: no parece pedido de skill

    context_text: historial de usuario + mensaje actual (URLs/punto de turnos previos).
    """
    text = user_message or ""
    context = (context_text or text or "").strip() or text
    lowered = text.lower().strip()
    if not lowered and not context.strip():
        return {"action": "none", "confidence": 0.0}

    # Continuidad: el usuario corrige/cuestiona un resultado del mismo hilo.
    if is_result_challenge_or_correction(text):
        return {
            "action": "none",
            "confidence": 0.2,
            "reason": "thread_challenge",
        }

    open_task = extract_open_task(
        text, context_text=context, thread_state=thread_state
    )
    if is_asking_for_needed_data(text) and open_task:
        return {
            "action": "clarify",
            "confidence": 0.8,
            "reply": ask_inputs_for_open_task(context, thread_state=thread_state),
            "reason": "user_asked_needed_data",
        }

    # Confirmación explícita de descarga tras una aclaración previa.
    if re.search(
        r"(?:s[ií]|dale|ok|okay).{0,30}descarg"
        r"|descarg(?:á|a|ar).{0,40}skill"
        r"|gener(?:á|a|ar).{0,40}skill"
        r"|busc(?:á|a|ar).{0,40}skill",
        lowered,
    ):
        if is_site_crawl_request(context):
            found_crawl = find_local_skill(context, arguments)
            if found_crawl.get("found") and found_crawl.get("id") == "crawl_domain_links":
                return {
                    "action": "execute",
                    "confidence": 0.9,
                    "skill": found_crawl,
                    "reason": "crawl_after_skill_confirm",
                }
        effective = resolve_effective_remote_task(text, context)
        if has_actionable_remote_task(effective):
            return {
                "action": "download",
                "confidence": 0.9,
                "reason": "explicit_download_confirm",
            }
        return {
            "action": "clarify",
            "confidence": 0.4,
            "reply": (
                "Puedo descargar/generar la skill, pero necesito la tarea real: "
                "qué dato querés, de qué URL/API y (si aplica) el código de punto. "
                "Ej.: «altura y caudal del punto 10009 en "
                "https://serviciosweb.cloud.irrigacion.gov.ar/public/telemetriaMovil»."
            ),
            "reason": "download_confirm_without_task",
        }

    if is_site_crawl_request(text) or is_site_crawl_request(context):
        found_crawl = find_local_skill(text if is_site_crawl_request(text) else context, arguments)
        if not found_crawl.get("found"):
            found_crawl = find_local_skill(
                "escanear subenlaces crawler " + (text or context),
                arguments,
            )
        if found_crawl.get("found"):
            args = dict(found_crawl.get("arguments") or {})
            missing = missing_fields_for_skill(str(found_crawl.get("id") or ""), args)
            if missing:
                return {
                    "action": "clarify",
                    "confidence": 0.7,
                    "skill": found_crawl,
                    "missing": missing,
                    "reply": clarifying_question_for_skill(found_crawl, missing),
                    "reason": "crawl_missing_url",
                }
            return {
                "action": "execute",
                "confidence": 0.92,
                "skill": found_crawl,
                "reason": "local_crawl_skill",
            }
        return {
            "action": "download",
            "confidence": 0.75,
            "reason": "crawl_no_local_fallback",
        }

    # Solo una URL (o URL + poco texto) después de pedir datos: usar contexto.
    if contains_url(text) and len(_tokenize(text)) <= 6:
        ctx_args = extract_web_skill_args(context)
        if is_telemetria_request(context) or ctx_args.get("punto") or any(
            k in context.lower()
            for k in ("altura", "caudal", "telemetr", "nivel", "medición", "medicion")
        ):
            return {
                "action": "download",
                "confidence": 0.85,
                "reason": "url_followup_with_context",
            }

    if not (
        is_action_request(text)
        or should_try_skill_marketplace(text)
        or looks_like_web_or_external_request(context)
        or is_thread_followup(text, context, thread_state)
    ):
        return {"action": "none", "confidence": 0.0}

    # Web/URL/telemetría: según la TAREA ABIERTA, no por keywords viejas del chat.
    web_probe = open_task or text
    if (
        looks_like_web_or_external_request(text)
        or looks_like_web_or_external_request(web_probe)
        or is_telemetria_request(web_probe)
    ):
        probe = (
            context
            if (open_task and (open_task in context)) or len(context) >= len(text)
            else text
        )
        missing = missing_web_clarify_fields(probe)
        if not missing:
            return {
                "action": "download",
                "confidence": 0.85,
                "reason": "web_request_clear",
            }
        # Telemetría con punto aunque falte URL explícita en este turno.
        if is_telemetria_request(probe) and "punto" not in missing:
            return {
                "action": "download",
                "confidence": 0.8,
                "reason": "telemetria_punto_clear",
            }
        return {
            "action": "clarify",
            "confidence": 0.4,
            "reply": clarifying_question_for_unknown(
                text, context_text=probe, thread_state=thread_state
            ),
            "reason": "web_request_unclear",
        }

    found = find_local_skill(text, arguments)
    if found.get("ambiguous"):
        return {
            "action": "clarify",
            "confidence": 0.45,
            "reply": clarifying_question_for_skill(
                {},
                [],
                ambiguous_candidates=found.get("candidates") or [],
            ),
            "reason": "ambiguous_catalog",
            "candidates": found.get("candidates"),
        }

    if found.get("found"):
        args = dict(found.get("arguments") or {})
        missing = missing_fields_for_skill(str(found.get("id") or ""), args)
        # Word: el contenido se genera; no pedir args numéricos.
        if str(found.get("id")) == "generar_documento_word":
            missing = []
        if missing:
            return {
                "action": "clarify",
                "confidence": 0.55,
                "skill": found,
                "missing": missing,
                "reply": clarifying_question_for_skill(found, missing),
                "reason": "missing_arguments",
            }
        conf = 0.9 if found.get("confidence") == "high" else 0.7
        return {
            "action": "execute",
            "confidence": conf,
            "skill": found,
            "reason": "local_match",
        }

    # Follow-up de la tarea ya abierta: continuar (no resetear al catálogo).
    if is_thread_followup(text, context, thread_state) and open_task:
        if has_concrete_task_inputs(text) or has_concrete_task_inputs(context):
            return {
                "action": "download",
                "confidence": 0.8,
                "reason": "followup_continue_open_task",
            }
        return {
            "action": "clarify",
            "confidence": 0.7,
            "reply": ask_inputs_for_open_task(
                context, thread_state=thread_state
            ),
            "reason": "followup_missing_inputs",
        }

    # Acción clara sin match: ofrecer descarga, no inventar skill local.
    if is_action_request(text) or should_try_skill_marketplace(text):
        # Si es muy vago ("podés hacer algo?"), preguntar.
        tokens = _tokenize(text)
        if len(tokens) < 4 and not re.search(r"\d", text):
            return {
                "action": "clarify",
                "confidence": 0.35,
                "reply": clarifying_question_for_unknown(
                    text, context_text=context, thread_state=thread_state
                ),
                "reason": "vague_action",
            }
        return {
            "action": "download",
            "confidence": 0.65,
            "reason": found.get("reason") or "no_local_skill",
        }

    return {"action": "none", "confidence": 0.0}


def download_remote_prompt() -> str:
    return (
        "No tengo una habilidad instalada para hacer eso todavía. "
        "¿Querés que descargue/genere la skill desde internet y la ejecute "
        "(con auditoría de Gemini)?"
    )


_DOWNLOAD_CONFIRM_ONLY_RE = re.compile(
    r"^\s*(?:s[ií]|dale|ok|okay|bueno|claro|perfecto)?[,.\s!]*"
    r"(?:por\s+favor[,.\s]*)?"
    r"(?:descarg(?:á|a|ar)|gener(?:á|a|ar)|busc(?:á|a|ar)|instal(?:á|a|ar))\w*"
    r".{0,40}\bskills?\b"
    r"[.\s!]*$",
    re.I,
)


def is_download_confirmation_only(text: str) -> bool:
    """True si el mensaje solo confirma descargar/generar skill (sin la tarea real)."""
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    if not cleaned or len(cleaned) > 120:
        return False
    if contains_url(cleaned):
        return False
    if re.search(
        r"\b(?:altura|caudal|punto|url|api|telemetr|web|http|consultar|entrar)\b",
        cleaned,
        re.I,
    ):
        return False
    return bool(_DOWNLOAD_CONFIRM_ONLY_RE.match(cleaned))


def resolve_effective_remote_task(
    user_message: str,
    conversation_context: str | None = None,
) -> str:
    """
    Tarea real para generar/ejecutar skills remotas.
    Si el usuario solo dice 'sí, descargá la skill', usa el historial.
    """
    task = (user_message or "").strip()
    context = (conversation_context or "").strip()
    if not context:
        return task

    parts = [p.strip() for p in re.split(r"\n\s*\n", context) if p.strip()]
    useful = [p for p in parts if not is_download_confirmation_only(p)]

    if is_download_confirmation_only(task):
        if useful:
            return "\n\n".join(useful)
        return context

    # Mensaje corto (URL sola, etc.): enriquecer con historial previo.
    if useful and task:
        if task not in useful[-1]:
            # Evitar duplicar si conversation_context ya incluye el mensaje actual.
            prior = [p for p in useful if p != task]
            if prior and (contains_url(task) or len(task) < 80):
                return "\n\n".join(prior + [task])
        return "\n\n".join(useful) if len("\n\n".join(useful)) >= len(task) else task

    return context if len(context) > len(task) else task


def has_actionable_remote_task(text: str) -> bool:
    """Hay sustancia para armar una skill (no solo 'descargá la skill')."""
    if not (text or "").strip():
        return False
    if is_download_confirmation_only(text):
        return False
    if contains_url(text) or is_telemetria_request(text):
        return True
    if extract_web_skill_args(text).get("punto"):
        return True
    lowered = text.lower()
    if any(
        k in lowered
        for k in (
            "altura",
            "caudal",
            "consultar",
            "entrar",
            "obtener",
            "traer",
            "scrap",
            "api",
            "word",
            "documento",
            "calcular",
            "análisis de red",
            "analisis de red",
            "wake-on-lan",
            "wake on lan",
        )
    ):
        return True
    # Texto con algo más que confirmación genérica.
    tokens = _tokenize(text)
    return len(tokens) >= 6


def conversation_context_text(
    user_message: str,
    history: list[dict[str, Any]] | None = None,
) -> str:
    """Une mensajes de usuario del historial + el actual (para URL/punto/tarea)."""
    parts: list[str] = []
    for item in history or []:
        role = (item.get("role") or "").lower()
        msg = (item.get("message") or "").strip()
        if role == "user" and msg:
            parts.append(msg)
    current = (user_message or "").strip()
    if current and (not parts or parts[-1] != current):
        parts.append(current)
    return "\n\n".join(parts)


def is_telemetria_request(text: str) -> bool:
    from app.services.skill_http import extract_urls

    lowered = (text or "").lower()
    if not lowered.strip():
        return False
    if is_site_crawl_request(text):
        return False
    if _TELEMETRY_INTENT_RE.search(lowered):
        return True
    if "dato-medicions" in lowered or "fulldto" in lowered:
        return True
    return any("telemetr" in u.lower() for u in extract_urls(text))


def extract_web_skill_args(text: str) -> dict[str, Any]:
    """Extrae url(s) y código de punto desde texto (mensaje o historial)."""
    from app.services.skill_http import extract_urls

    args: dict[str, Any] = {}
    urls = extract_urls(text)
    if urls:
        args["urls"] = urls
        args["url"] = urls[0]
    punto = re.search(
        r"(?:punto|estaci[oó]n|c[oó]digo|codigo|id)\s*[:#-]?\s*(\d{3,8})",
        text or "",
        re.I,
    )
    if punto:
        args["punto"] = punto.group(1)
    elif is_telemetria_request(text) and re.search(r"\b\d{4,6}\b", text or ""):
        # Evitar tomar años tipo 2026 si hay contexto de telemetría/URL.
        candidates = re.findall(r"\b(\d{4,6})\b", text or "")
        for cand in candidates:
            if cand.startswith("20") and len(cand) == 4:
                continue
            args["punto"] = cand
            break
    return args


def merge_web_skill_arguments(
    arguments: dict[str, Any] | None,
    text: str,
) -> dict[str, Any]:
    """Completa url/punto/api_url para skills web/telemetría sin pisar valores útiles."""
    merged = dict(arguments or {})
    extracted = extract_web_skill_args(text)
    for key, value in extracted.items():
        if merged.get(key) in (None, "", [], {}):
            merged[key] = value
    if is_telemetria_request(text) or any(
        "telemetr" in str(u).lower() for u in (merged.get("urls") or [])
    ):
        merged.setdefault(
            "api_url",
            "https://serviciosweb.cloud.irrigacion.gov.ar/services/telemetria/"
            "api/public/dato-medicions/fullDto",
        )
    if text and merged.get("query") in (None, "", {}):
        merged["query"] = text
    return merged


def _extract_numbers(text: str) -> list[float]:
    """Extrae cantidades; ignora dígitos que son parte de unidades (m3, m2, etc.)."""
    cleaned = text or ""
    cleaned = re.sub(
        r"(?i)(?:\b(?:ha|cm|mm|l/?s|lps|m3/s|m³/s|m3/h|m3/d)\b"
        r"|(?<=\d)\s*m[²³23]\b"
        r"|\bm[²³23]\b"
        r"|m\^[23]|m³|m²)",
        " ",
        cleaned,
    )
    raw = re.findall(r"(\d+(?:[.,]\d+)?)", cleaned)
    out: list[float] = []
    for item in raw:
        try:
            out.append(float(item.replace(",", ".")))
        except ValueError:
            continue
    return out


def infer_arguments(skill_id: str, user_message: str) -> dict[str, Any]:
    """Inferencia heurística de argumentos cuando el LLM no los extrajo."""
    text = user_message or ""
    lowered = text.lower()
    numbers = _extract_numbers(text)
    args: dict[str, Any] = {"query": text}

    if skill_id == "caudal_canal":
        if len(numbers) >= 2:
            args["area_m2"] = numbers[0]
            args["velocidad_ms"] = numbers[1]
        elif len(numbers) == 1:
            if any(k in lowered for k in ("area", "área", "m2", "m²", "seccion", "sección")):
                args["area_m2"] = numbers[0]
            else:
                args["velocidad_ms"] = numbers[0]
    elif skill_id == "conversion_unidades":
        if numbers:
            args["valor"] = numbers[0]
        for unit in ("l/s", "l/s", "m3/s", "m³/s", "m3/h", "m3/d", "m3/dia", "m3/día"):
            if unit.replace("³", "3") in lowered.replace("³", "3"):
                args["unidad"] = unit
                break
    elif skill_id == "prorrateo_turno":
        if numbers:
            args["total"] = numbers[0]
        pesos = re.findall(r"(\d+(?:[.,]\d+)?)\s*(?:ha|acciones?|derechos?)", lowered)
        if pesos:
            args["partes"] = [{"nombre": f"parte_{i+1}", "peso": float(p.replace(",", "."))} for i, p in enumerate(pesos)]
    elif skill_id == "lamina_riego":
        # Solo asignar volumen si el texto habla de volumen/lámina (no un "punto 10009").
        if any(k in lowered for k in ("volumen", "m3", "m³", "lamina", "lámina", "mm")):
            if numbers:
                args["volumen_m3"] = numbers[0]
                if len(numbers) > 1:
                    if re.search(r"\bha\b|hect", lowered):
                        args["superficie_ha"] = numbers[1]
                    elif re.search(r"m[²2]|metro", lowered):
                        args["superficie_m2"] = numbers[1]
                    # Sin unidad de superficie → no adivinar el 2.º número.
        elif re.search(r"\bha\b|hect", lowered) and numbers:
            args["superficie_ha"] = numbers[0]
            if len(numbers) > 1:
                args["volumen_m3"] = numbers[1]
    elif skill_id == "tiempo_riego":
        if any(k in lowered for k in ("volumen", "m3", "m³", "tiempo", "regar", "caudal")):
            if numbers:
                args["volumen_m3"] = numbers[0]
                if len(numbers) > 1:
                    if "l/s" in lowered or "lps" in lowered or "litro" in lowered:
                        args["caudal_ls"] = numbers[1]
                    elif re.search(r"m[³3]/s|caudal", lowered):
                        args["caudal_m3s"] = numbers[1]
    elif skill_id == "generar_documento_word":
        title_match = re.search(
            r"(?:titulo|título)\s*[:\-]\s*(.+)$", text, re.I | re.M
        )
        if title_match:
            args["titulo"] = title_match.group(1).strip()[:120]
        # No copiar el pedido del usuario como "contenido": eso se redacta en
        # prepare_docx_arguments() con el LLM.
        args["pedido"] = text
    elif skill_id == "crawl_domain_links":
        from app.services.skill_http import extract_urls

        urls = extract_urls(text)
        if urls:
            args["base_url"] = urls[0]
            args["url"] = urls[0]
        keys = [
            k
            for k in ("normativa", "resolucion", "resolución", "ley", "institucional", "pdf")
            if k in lowered
        ]
        if keys:
            args["keywords"] = keys
    return args


def enrich_skill_arguments(skill: dict[str, Any], user_message: str) -> dict[str, Any]:
    """Completa argumentos faltantes antes de ejecutar la skill."""
    skill_id = str(skill.get("id") or "")
    current = dict(skill.get("arguments") or {})
    inferred = infer_arguments(skill_id, user_message)
    merged = {**inferred, **{k: v for k, v in current.items() if v not in (None, "", {})}}
    if skill_id == "generar_documento_word":
        if not merged.get("titulo") and not merged.get("title"):
            merged["titulo"] = "Documento Irrigación"
        merged["pedido"] = user_message
    if skill_id == "crawl_domain_links":
        if not merged.get("base_url") and merged.get("url"):
            merged["base_url"] = merged["url"]
        if not merged.get("url") and merged.get("base_url"):
            merged["url"] = merged["base_url"]
    return merged


def prepare_docx_arguments(user_message: str, arguments: dict[str, Any]) -> dict[str, Any]:
    """Redacta título + bloques con formato real del Word (no el texto del pedido)."""
    from langchain_core.messages import HumanMessage, SystemMessage
    from langchain_openai import ChatOpenAI

    from app.core.config import get_settings
    from app.services.token_guard import fit_user_message

    merged = dict(arguments or {})
    existing_blocks = merged.get("blocks") or merged.get("bloques")
    if isinstance(existing_blocks, list) and existing_blocks:
        return merged

    settings = get_settings()
    llm = ChatOpenAI(
        model=settings.chat_model,
        api_key=settings.groq_api_key,
        base_url=settings.groq_base_url,
        temperature=0.55,
    )
    pedido = fit_user_message(user_message or merged.get("pedido") or "")
    prompt = """
Sos un redactor técnico de Irrigación de Malargüe. El usuario pidió un documento Word.
Generá el CONTENIDO REAL del documento (nunca copies el pedido literal) y aplicá el FORMATO
que pidió (títulos, subtítulos, viñetas, numeración, tablas, negritas, etc.).

Reglas:
- Si pide N párrafos aleatorios: inventá exactamente N párrafos (type=paragraph).
- Si pide secciones/títulos: usá type=heading con level 1..3.
- Si pide lista con viñetas: type=bullet + items[].
- Si pide lista numerada: type=number + items[].
- Si pide tabla o datos tabulares: type=table + rows[][] (primera fila = encabezados si aplica).
- Negrita con **texto** y cursiva con *texto* dentro de text/items.
- Español rioplatense, tono institucional claro.
- Respondé ÚNICAMENTE JSON válido (sin markdown) con esta forma:
{
  "titulo": "Título del documento",
  "filename": "opcional.docx",
  "title_align": "center",
  "blocks": [
    {"type": "heading", "level": 1, "text": "Sección"},
    {"type": "paragraph", "text": "Párrafo con **negrita** si hace falta", "align": "justify"},
    {"type": "bullet", "items": ["ítem 1", "ítem 2"]},
    {"type": "number", "items": ["paso 1", "paso 2"]},
    {"type": "table", "rows": [["Col A", "Col B"], ["v1", "v2"]]}
  ]
}
""".strip()
    try:
        response = llm.invoke(
            [
                SystemMessage(content=prompt),
                HumanMessage(content=f"Pedido del usuario:\n{pedido}"),
            ]
        )
        raw = (getattr(response, "content", None) or "").strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
        data = json.loads(raw)
        if isinstance(data, dict):
            title = str(data.get("titulo") or data.get("title") or "").strip()
            if title:
                merged["titulo"] = title[:120]
            filename = str(data.get("filename") or data.get("nombre") or "").strip()
            if filename:
                merged["filename"] = filename
            if data.get("title_align"):
                merged["title_align"] = data.get("title_align")
            blocks = data.get("blocks") or data.get("bloques")
            if isinstance(blocks, list) and blocks:
                merged["blocks"] = blocks
            else:
                content = str(data.get("contenido") or data.get("content") or "").strip()
                paragraphs = data.get("paragraphs") or data.get("parrafos")
                if isinstance(paragraphs, list) and paragraphs:
                    merged["blocks"] = [
                        {"type": "paragraph", "text": str(p).strip()}
                        for p in paragraphs
                        if str(p).strip()
                    ]
                elif content:
                    parts = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]
                    merged["blocks"] = [{"type": "paragraph", "text": p} for p in parts]
                    merged["contenido"] = content
    except Exception:
        n_match = re.search(r"(\d+)\s*p[aá]rrafos?", (user_message or "").lower())
        n = int(n_match.group(1)) if n_match else 3
        n = max(1, min(n, 12))
        stubs = [
            (
                "En la cuenca del río Atuel, el control de caudales requiere "
                "mediciones periódicas en tomas y canales secundarios."
            ),
            (
                "El prorrateo de turnos de riego se realiza según derechos "
                "consuntivos y superficies empadronadas en cada zona."
            ),
            (
                "La documentación institucional debe registrar volúmenes "
                "entregados, incidencias operativas y acuerdos de distribución."
            ),
            (
                "Las láminas de riego se estiman a partir del volumen aplicado "
                "y la superficie efectiva bajo riego."
            ),
            (
                "Un informe técnico claro facilita la coordinación entre "
                "inspectores de cauce y usuarios de riego."
            ),
        ]
        merged.setdefault("titulo", "Documento Irrigación")
        merged["blocks"] = [
            {"type": "paragraph", "text": stubs[i % len(stubs)], "align": "justify"}
            for i in range(n)
        ]

    return merged


# Esquemas de argumentos esperados por skill (para extracción inteligente vía LLM).
SKILL_ARG_SCHEMAS: dict[str, dict[str, Any]] = {
    "caudal_canal": {
        "fields": {
            "area_m2": "Área de la sección en m²",
            "velocidad_ms": "Velocidad del agua en m/s",
        },
        "required": ["area_m2", "velocidad_ms"],
        "notes": "Q = A·v. Inferí área y velocidad aunque estén en otra unidad y convertí a m² y m/s.",
    },
    "conversion_unidades": {
        "fields": {
            "valor": "Número a convertir",
            "unidad": "Unidad de origen: l/s, m3/s, m3/h o m3/d",
        },
        "required": ["valor", "unidad"],
        "notes": "Normalizá la unidad a una de: l/s, m3/s, m3/h, m3/d.",
    },
    "prorrateo_turno": {
        "fields": {
            "total": "Caudal o volumen total a repartir",
            "partes": "Lista de {nombre, peso} (acciones, ha o derechos)",
        },
        "required": ["total", "partes"],
        "notes": "Si dice '3 usuarios con 2, 3 y 5 acciones' armá partes con esos pesos.",
    },
    "lamina_riego": {
        "fields": {
            "volumen_m3": "Volumen aplicado en m³",
            "superficie_ha": "Superficie en hectáreas (preferida)",
            "superficie_m2": "Superficie en m² (alternativa)",
        },
        "required": ["volumen_m3"],
        "notes": "Debe haber superficie_ha o superficie_m2.",
    },
    "tiempo_riego": {
        "fields": {
            "volumen_m3": "Volumen a aplicar en m³",
            "caudal_ls": "Caudal en L/s (preferido)",
            "caudal_m3s": "Caudal en m³/s (alternativa)",
        },
        "required": ["volumen_m3"],
        "notes": "Debe haber caudal_ls o caudal_m3s.",
    },
    "crawl_domain_links": {
        "fields": {
            "base_url": "URL del sitio a escanear (ej. https://www.irrigacion.gov.ar/web/)",
            "keywords": "Filtros opcionales (normativa, resolucion, ley, institucional, pdf)",
        },
        "required": ["base_url"],
        "notes": "No es telemetría: no pidas código de estación. Solo la URL base.",
    },
}


def _numeric_ready(value: Any) -> bool:
    if value in (None, "", {}, []):
        return False
    if isinstance(value, (int, float)):
        return True
    try:
        float(str(value).replace(",", "."))
        return True
    except (TypeError, ValueError):
        return False


def _args_complete_for_skill(skill_id: str, args: dict[str, Any]) -> bool:
    schema = SKILL_ARG_SCHEMAS.get(skill_id)
    if not schema:
        return bool(args)
    for key in schema.get("required") or []:
        if key == "partes":
            partes = args.get("partes") or args.get("derechos")
            if not isinstance(partes, list) or not partes:
                return False
            continue
        if key in {"base_url", "url"}:
            if not str(args.get("base_url") or args.get("url") or "").strip():
                return False
            continue
        if not _numeric_ready(args.get(key)):
            return False
    if skill_id == "lamina_riego":
        return _numeric_ready(args.get("superficie_ha")) or _numeric_ready(
            args.get("superficie_m2")
        )
    if skill_id == "tiempo_riego":
        return _numeric_ready(args.get("caudal_ls")) or _numeric_ready(
            args.get("caudal_m3s")
        )
    return True


def _llm_extract_skill_arguments(
    skill: dict[str, Any],
    user_message: str,
    current: dict[str, Any],
) -> dict[str, Any]:
    from langchain_core.messages import HumanMessage, SystemMessage
    from langchain_openai import ChatOpenAI

    from app.core.config import get_settings
    from app.services.token_guard import fit_user_message

    skill_id = str(skill.get("id") or "")
    schema = SKILL_ARG_SCHEMAS.get(skill_id) or {
        "fields": {},
        "required": [],
        "notes": skill.get("description") or "",
    }
    settings = get_settings()
    llm = ChatOpenAI(
        model=settings.chat_model,
        api_key=settings.groq_api_key,
        base_url=settings.groq_base_url,
        temperature=0.1,
    )
    prompt = f"""
Sos un extractor de argumentos para skills de Irrigación de Malargüe.
Skill: {skill.get('name') or skill_id}
Descripción: {skill.get('description') or ''}
Campos esperados: {json.dumps(schema.get('fields') or {}, ensure_ascii=False)}
Requeridos: {json.dumps(schema.get('required') or [], ensure_ascii=False)}
Notas: {schema.get('notes') or ''}

Extraé del pedido del usuario SOLO los argumentos de la skill.
- Inferí unidades y convertí a las pedidas cuando sea obvio (ha↔m², L/s↔m³/s).
- No inventes números que el usuario no dio.
- Si falta un dato imprescindible, igual devolvé lo que puedas y dejá ausentes los demás.
- Respondé ÚNICAMENTE JSON (objeto) con los campos.
""".strip()
    try:
        response = llm.invoke(
            [
                SystemMessage(content=prompt),
                HumanMessage(
                    content=(
                        f"Pedido:\n{fit_user_message(user_message)}\n\n"
                        f"Parcial ya inferido:\n{json.dumps(current, ensure_ascii=False)}"
                    )
                ),
            ]
        )
        raw = (getattr(response, "content", None) or "").strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
        data = json.loads(raw)
        if isinstance(data, dict):
            cleaned = {k: v for k, v in data.items() if v not in (None, "", [], {})}
            return {**current, **cleaned}
    except Exception:
        pass
    return current


def prepare_skill_arguments(
    skill: dict[str, Any],
    user_message: str,
    *,
    context_text: str | None = None,
) -> dict[str, Any]:
    """
    Personaliza argumentos de cualquier skill:
    - Word: redacta contenido + formato
    - Cálculos: completa con heurística y, si falta, extracción LLM
    - Remotas: completa url/punto desde historial; LLM solo si sigue vacío
    """
    skill_id = str(skill.get("id") or "")
    ctx = (context_text or user_message or "").strip()
    base = enrich_skill_arguments(skill, user_message)
    base = merge_web_skill_arguments(base, ctx)
    if ctx and (skill.get("source") == "remote" or skill_id.startswith("remote_")):
        base["query"] = ctx

    if skill_id == "generar_documento_word":
        return prepare_docx_arguments(user_message, base)

    if skill_id in SKILL_ARG_SCHEMAS:
        if _args_complete_for_skill(skill_id, base):
            return base
        return _llm_extract_skill_arguments(skill, user_message, base)

    # Skills remotas / desconocidas: si hay pocos args útiles, pedir al LLM.
    useful = {
        k: v
        for k, v in base.items()
        if k not in {"query", "pedido", "raw"} and v not in (None, "", {}, [])
    }
    if useful:
        return base
    return _llm_extract_skill_arguments(skill, ctx or user_message, base)


def detect_response_style(user_message: str) -> str:
    """Preferencia de presentación pedida por el usuario (chat o post-skill)."""
    lowered = (user_message or "").lower()
    hints: list[str] = []
    if any(k in lowered for k in ("tabla", "tabular", "markdown table", "|")):
        hints.append("Presentá números en tabla Markdown.")
    if any(k in lowered for k in ("viñeta", "vineta", "bullet", "lista", "ítems", "items")):
        hints.append("Usá lista con viñetas.")
    if any(k in lowered for k in ("paso a paso", "pasos", "numerad")):
        hints.append("Explicá en pasos numerados.")
    if any(k in lowered for k in ("breve", "corto", "resumí", "resumi", "tl;dr", "en una línea")):
        hints.append("Sé muy breve (máximo 5 líneas).")
    if any(k in lowered for k in ("detall", "explicá", "explica", "desarroll")):
        hints.append("Dale una explicación un poco más detallada, sin relleno.")
    if any(k in lowered for k in ("formal", "institucional", "para el expediente")):
        hints.append("Tono más formal/institucional.")
    if any(k in lowered for k in ("simple", "para un peón", "como si fuera", "en criollo")):
        hints.append("Lenguaje simple y directo, sin jerga innecesaria.")
    if not hints:
        if is_casual_chat(user_message):
            return (
                "Es charla. Contestá cálido y natural, 1-3 líneas, como un "
                "asistente de chat. Sin menú ni 'estoy funcionando'."
            )
        return (
            "Hablá como un asistente de chat (cálido, claro). "
            "Si hay varios números, usá tabla o viñetas."
        )
    return " ".join(hints)


def _tokenize(text: str) -> set[str]:
    return set(re.findall(r"[a-záéíóúüñ0-9/]+", (text or "").lower()))


def search_catalog(task: str, arguments: dict[str, Any] | None = None) -> dict[str, Any]:
    """Busca la skill más pertinente en el catálogo institucional."""
    ranked = rank_catalog_skills(task)
    available = [
        {"id": s["id"], "name": s["name"], "description": s["description"]}
        for s in CATALOG
    ]
    if not ranked or ranked[0][1] < 1:
        return {"found": False, "query": task, "available": available}
    best, best_score = ranked[0]
    resolved_args = arguments or {}
    if best["id"] == "generar_documento_word" or not any(
        k not in {"query", "raw", "valor"} for k in resolved_args
    ):
        resolved_args = {**infer_arguments(best["id"], task), **resolved_args}
    return {
        "found": True,
        "id": best["id"],
        "name": best["name"],
        "description": best["description"],
        "code": best["code"],
        "arguments": resolved_args,
        "score": best_score,
    }


def _parse_arguments(raw: Any) -> dict[str, Any]:
    if raw is None or raw == "":
        return {}
    if isinstance(raw, dict):
        return raw
    if isinstance(raw, str):
        try:
            data = json.loads(raw)
            return data if isinstance(data, dict) else {"valor": data}
        except json.JSONDecodeError:
            return {"raw": raw}
    return {"valor": raw}


@tool
def search_skill_marketplace(task: str, arguments_json: str = "{}") -> str:
    """Busca en el marketplace local/catálogo de skills una herramienta o script Python
    ejecutable para realizar tareas técnicas solicitadas por el usuario.

    Usala cuando el usuario pide un cálculo, conversión, prorrateo, lámina o tiempo de riego,
    generación de documentos Word (con formato), test de red/velocidad, o cualquier
    automatización / "buscar una skill". NUNCA interpretes "skill" como Alexa, Google
    Assistant, Siri o Cortana. No inventes números: primero buscá la skill.

    Args:
        task: descripción breve de la tarea (ej. 'calcular caudal con área y velocidad').
        arguments_json: JSON con números/datos/unidades extraídos del mensaje
            (ej. area_m2, velocidad_ms, valor, unidad, partes, volumen_m3, superficie_ha).
    """
    result = search_catalog(task, _parse_arguments(arguments_json))
    public = {k: v for k, v in result.items() if k != "code"}
    if result.get("found"):
        public["has_code"] = True
    return json.dumps(public, ensure_ascii=False)
