"""Ingestión de documentos: extracción de texto, chunking, embeddings y persistencia."""

from __future__ import annotations

import logging
import tempfile
from pathlib import Path

import pymupdf as fitz
from docx import Document
from google.genai import types
from sqlalchemy import text
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.gemini import gemini_client

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg"}

OCR_SYSTEM_PROMPT = (
    "Eres un motor OCR estricto. Transcribe TODO el texto visible en la imagen "
    "a Markdown. Respeta tablas (usa sintaxis Markdown de tablas), encabezados, "
    "listas y datos numéricos. No inventes contenido. No agregues comentarios "
    "ni explicaciones: solo el texto transcripto."
)


def detect_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Tipo de archivo no soportado: '{ext}'. "
            f"Permitidos: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return ext


def split_text(text_content: str, chunk_size: int | None = None, overlap: int | None = None) -> list[str]:
    settings = get_settings()
    chunk_size = chunk_size or settings.chunk_size
    overlap = overlap or settings.chunk_overlap

    cleaned = text_content.strip()
    if not cleaned:
        return []

    if len(cleaned) <= chunk_size:
        return [cleaned]

    chunks: list[str] = []
    start = 0
    text_len = len(cleaned)

    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = cleaned[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= text_len:
            break
        start = max(0, end - overlap)

    return chunks


def generate_embeddings(
    texts: list[str],
    *,
    task_type: str = "RETRIEVAL_DOCUMENT",
) -> list[list[float]]:
    """Embeddings con Gemini (gemini-embedding-001 → 768 dims)."""
    if not texts:
        return []

    settings = get_settings()
    client = gemini_client()
    vectors: list[list[float]] = []

    # Gemini embed_content acepta un contenido por llamada; batched simple.
    for piece in texts:
        response = client.models.embed_content(
            model=settings.embedding_model,
            contents=piece,
            config={
                "output_dimensionality": settings.embedding_dimensions,
                "task_type": task_type,
            },
        )
        # google-genai: response.embeddings[0].values ó response.embedding.values
        embedding = _extract_embedding_values(response)
        if len(embedding) != settings.embedding_dimensions:
            logger.warning(
                "Dimensión de embedding inesperada: %s (esperado %s)",
                len(embedding),
                settings.embedding_dimensions,
            )
        vectors.append(embedding)

    return vectors


def _extract_embedding_values(response: object) -> list[float]:
    embeddings = getattr(response, "embeddings", None)
    if embeddings:
        first = embeddings[0]
        values = getattr(first, "values", None)
        if values is not None:
            return [float(v) for v in values]
    single = getattr(response, "embedding", None)
    if single is not None:
        values = getattr(single, "values", None)
        if values is not None:
            return [float(v) for v in values]
    raise RuntimeError("Respuesta de embedding Gemini sin values")


def _ocr_image_bytes(image_bytes: bytes, mime: str = "image/png") -> str:
    """OCR multimodal con Gemini (gratis / flash)."""
    settings = get_settings()
    client = gemini_client()
    response = client.models.generate_content(
        model=settings.ocr_model,
        contents=[
            types.Part.from_bytes(data=image_bytes, mime_type=mime),
            f"{OCR_SYSTEM_PROMPT}\n\nTranscribe esta imagen a Markdown.",
        ],
        config=types.GenerateContentConfig(temperature=0),
    )
    return (response.text or "").strip()


def _extract_text_from_pdf(path: Path) -> str:
    settings = get_settings()
    doc = fitz.open(path)
    try:
        native_parts: list[str] = []
        for page in doc:
            native_parts.append(page.get_text("text") or "")

        native_text = "\n".join(native_parts).strip()
        if len(native_text) >= settings.scanned_pdf_char_threshold:
            return native_text

        logger.info(
            "PDF escaneado o con poco texto (%s chars). Usando OCR Gemini (%s).",
            len(native_text),
            settings.ocr_model,
        )
        ocr_parts: list[str] = []
        for page_index, page in enumerate(doc):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            image_bytes = pix.tobytes("png")
            page_text = _ocr_image_bytes(image_bytes, mime="image/png")
            if page_text:
                ocr_parts.append(f"## Página {page_index + 1}\n\n{page_text}")
        return "\n\n".join(ocr_parts).strip()
    finally:
        doc.close()


def _extract_text_from_docx(path: Path) -> str:
    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_blocks: list[str] = []
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            table_blocks.append("\n".join(rows))

    parts = paragraphs + table_blocks
    return "\n\n".join(parts).strip()


def _extract_text_from_image(path: Path) -> str:
    ext = path.suffix.lower()
    mime = "image/jpeg" if ext in {".jpg", ".jpeg"} else "image/png"
    return _ocr_image_bytes(path.read_bytes(), mime=mime)


def extract_text(path: Path, filename: str) -> str:
    ext = detect_file_type(filename)
    if ext == ".pdf":
        return _extract_text_from_pdf(path)
    if ext == ".docx":
        return _extract_text_from_docx(path)
    if ext in {".png", ".jpg", ".jpeg"}:
        return _extract_text_from_image(path)
    raise ValueError(f"Extensión no manejada: {ext}")


def persist_chunks(
    db: Session,
    document_name: str,
    chunks: list[str],
    embeddings: list[list[float]],
    *,
    scope: str = "irrigacion",
    user_id: str | None = None,
    source: str = "upload",
    title: str | None = None,
) -> int:
    if len(chunks) != len(embeddings):
        raise ValueError("La cantidad de chunks y embeddings no coincide")

    scope_norm = (scope or "irrigacion").strip().lower()
    if scope_norm not in {"irrigacion", "personal"}:
        raise ValueError("scope debe ser 'irrigacion' o 'personal'")
    if scope_norm == "personal" and not user_id:
        raise ValueError("user_id es obligatorio para contexto personal")

    insert_sql = text(
        """
        INSERT INTO document_chunks (
            document_name, content, embedding, scope, user_id, source, title
        )
        VALUES (
            :document_name, :content, CAST(:embedding AS vector),
            :scope, CAST(:user_id AS uuid), :source, :title
        )
        """
    )

    for chunk, embedding in zip(chunks, embeddings, strict=True):
        embedding_literal = "[" + ",".join(str(float(v)) for v in embedding) + "]"
        db.execute(
            insert_sql,
            {
                "document_name": document_name,
                "content": chunk,
                "embedding": embedding_literal,
                "scope": scope_norm,
                "user_id": user_id,
                "source": source or "upload",
                "title": title,
            },
        )
    db.commit()
    return len(chunks)


def ingest_text_note(
    db: Session,
    content: str,
    *,
    scope: str,
    user_id: str | None = None,
    title: str | None = None,
    document_name: str | None = None,
) -> dict:
    """Indexa una nota de chat como contexto tipado (personal / irrigación)."""
    cleaned = (content or "").strip()
    if not cleaned:
        raise ValueError("No hay texto para guardar como contexto")

    scope_norm = (scope or "").strip().lower()
    if scope_norm in {"irrigacion", "de irrigacion", "oficina", "institucional"}:
        scope_norm = "irrigacion"
    elif scope_norm in {"personal", "mio", "mío", "privado"}:
        scope_norm = "personal"
    if scope_norm not in {"irrigacion", "personal"}:
        raise ValueError("scope debe ser 'irrigacion' o 'personal'")
    if scope_norm == "personal" and not user_id:
        raise ValueError(
            "Para guardar contexto personal necesitás iniciar sesión con Google."
        )

    label = title or cleaned.split("\n", 1)[0][:80]
    doc_name = document_name or (
        f"nota-personal:{label}" if scope_norm == "personal" else f"nota-irrigacion:{label}"
    )
    chunks = split_text(cleaned)
    if not chunks:
        chunks = [cleaned]
    embeddings = generate_embeddings(chunks)
    created = persist_chunks(
        db,
        doc_name,
        chunks,
        embeddings,
        scope=scope_norm,
        user_id=user_id,
        source="chat_note",
        title=label,
    )
    return {
        "chunks_created": created,
        "scope": scope_norm,
        "title": label,
        "document_name": doc_name,
    }


def ingest_file(
    db: Session,
    file_bytes: bytes,
    filename: str,
    *,
    scope: str = "irrigacion",
    user_id: str | None = None,
) -> dict:
    """Procesa un archivo subido y lo indexa en document_chunks."""
    detect_file_type(filename)
    suffix = Path(filename).suffix.lower()

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=True) as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        tmp_path = Path(tmp.name)

        extracted = extract_text(tmp_path, filename)
        if not extracted.strip():
            return {
                "filename": filename,
                "chunks_created": 0,
                "warning": "No se pudo extraer texto del archivo",
            }

        chunks = split_text(extracted)
        if not chunks:
            return {
                "filename": filename,
                "chunks_created": 0,
                "warning": "El texto extraído no produjo chunks",
            }

        embeddings = generate_embeddings(chunks)
        created = persist_chunks(
            db,
            filename,
            chunks,
            embeddings,
            scope=scope,
            user_id=user_id,
            source="upload",
            title=filename,
        )

        return {
            "filename": filename,
            "chunks_created": created,
            "scope": scope,
        }

